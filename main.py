import numpy as np # linear algebra
import pandas as pd # data processing, CSV file I/O (e.g. pd.read_csv)
import spacy
import pickle
import random
import sys, fitz
import docx

def parse_resume(file_path):
    text= get_text(file_path)
    segment = classify_resume(text)
    return segment

#Extract text from DOCX
def get_text_doc(file_path):
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text
    print(text)
    return text

# Get text from pdf
def get_text_pdf(file_path):
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text = text + str(page.getText())
    txt = " ".join(text.split('\n'))
    return txt


def get_text(file_path):
    """Get text from pdf or word"""
    print(file_path[-3:])
    if file_path[-3:]=="pdf":
        txt = get_text_pdf(file_path)
        
    elif file_path[-3:]=="doc" or file_path[-4:]=="docx":
        txt=get_text_doc(file_path)
        # doc = docx.Document(file_path)
        # text = ""

        # for line in doc.paragraphs:
        #     text += line.text
        # txt = " ".join(text.split('\n'))
    return txt

def classify_resume(txt):
    """
    Classifying the resume using NER
    """
    nlp_model = spacy.load('nlp_ner_model')
    # Applying the model
    doc = nlp_model(txt)
    segment = ""
    for ent in doc.ents:
        segment += (f'{ent.label_.upper():{30}}- {ent.text}\n')
    return segment

print(parse_resume("data/Chang_Ker_Fui.docx"))